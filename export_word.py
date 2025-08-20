from docx import Document
from docx.shared import Inches
from datetime import datetime

def export_to_word(df, indicator, units, results, graph_path):
    doc = Document()

    # Заголовок
    doc.add_heading(f"Показник: {indicator} ({units})", level=1)

    # Таблиця з сирими даними
    doc.add_heading("Початкові дані", level=2)
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    table.style = "Table Grid"

    # Заголовки
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)

    # Дані
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = str(df.iat[i, j]) if df.iat[i, j] is not None else ""
            table.cell(i+1, j).text = val

    # Результати аналізів
    doc.add_heading("Результати аналізу", level=2)
    for key, value in results.items():
        doc.add_paragraph(f"{key}: {value}")

    # Додаємо графік
    doc.add_heading("Графічне відображення", level=2)
    doc.add_picture(graph_path, width=Inches(5))

    # Дата і назва програми
    doc.add_paragraph(f"\nДата виконання: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("Програма: SAD - Статистичний аналіз даних")

    doc.save("Результати_аналізу.docx")
