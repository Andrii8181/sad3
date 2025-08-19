from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QCheckBox, QLineEdit, QMessageBox
import pandas as pd
from scipy import stats
import statsmodels.api as sm
import matplotlib.pyplot as plt
import docx
from datetime import datetime


class AnalysisWindow(QDialog):
    def __init__(self, df):
        super().__init__()
        self.df = df
        self.setWindowTitle("Аналіз даних")
        self.setGeometry(250, 150, 400, 300)

        layout = QVBoxLayout()

        # Поля для назви показника і одиниць виміру
        layout.addWidget(QLabel("Назва показника:"))
        self.indicator_input = QLineEdit()
        layout.addWidget(self.indicator_input)

        layout.addWidget(QLabel("Одиниці виміру:"))
        self.unit_input = QLineEdit()
        layout.addWidget(self.unit_input)

        # Чекбокси для вибору аналізів
        self.chk_anova = QCheckBox("Однофакторний дисперсійний аналіз (ANOVA)")
        self.chk_ttest = QCheckBox("t-тест (Student)")
        self.chk_median = QCheckBox("Медіанний тест (непараметричний)")
        self.chk_graph = QCheckBox("Побудувати графік (середні значення)")

        layout.addWidget(self.chk_anova)
        layout.addWidget(self.chk_ttest)
        layout.addWidget(self.chk_median)
        layout.addWidget(self.chk_graph)

        # Кнопка запуску
        self.run_button = QPushButton("Виконати аналіз")
        self.run_button.clicked.connect(self.run_analysis)
        layout.addWidget(self.run_button)

        self.setLayout(layout)

    def run_analysis(self):
        results = []
        numeric_df = self.df.apply(pd.to_numeric, errors="coerce").dropna()

        # Перевірка на нормальність Шапіро–Вілка
        w, pvalue = stats.shapiro(numeric_df.iloc[:, -1])
        if pvalue < 0.05:
            QMessageBox.warning(self, "Перевірка нормальності", 
                                "Дані не відповідають нормальному розподілу!\n"
                                "Рекомендується використовувати непараметричні методи.")
        else:
            QMessageBox.information(self, "Перевірка нормальності", "Дані відповідають нормальному розподілу.")

        # Однофакторний ANOVA
        if self.chk_anova.isChecked():
            try:
                groups = [numeric_df.iloc[:, -1]]
                f_val, p_val = stats.f_oneway(*groups)
                results.append(f"ANOVA: F = {f_val:.3f}, p = {p_val:.4f}")
            except Exception as e:
                results.append(f"ANOVA: помилка ({str(e)})")

        # T-тест
        if self.chk_ttest.isChecked():
            try:
                sample1 = numeric_df.iloc[:, 0]
                sample2 = numeric_df.iloc[:, 1]
                t_val, p_val = stats.ttest_ind(sample1, sample2)
                results.append(f"T-тест: t = {t_val:.3f}, p = {p_val:.4f}")
            except Exception as e:
                results.append(f"T-тест: помилка ({str(e)})")

        # Медіанний тест
        if self.chk_median.isChecked():
            try:
                stat, p_val, med, tbl = stats.median_test(*[numeric_df[col] for col in numeric_df.columns])
                results.append(f"Медіанний тест: χ² = {stat:.3f}, p = {p_val:.4f}")
            except Exception as e:
                results.append(f"Медіанний тест: помилка ({str(e)})")

        # Графік
        if self.chk_graph.isChecked():
            plt.figure()
            numeric_df.mean().plot(kind="bar")
            plt.title("Середні значення по факторах")
            plt.savefig("results_plot.png")
            plt.close()
            results.append("Побудовано графік (results_plot.png)")

        # Запис у Word
        doc = docx.Document()
        doc.add_heading(f"Аналіз даних – {self.indicator_input.text()}", 0)
        doc.add_paragraph(f"Одиниці виміру: {self.unit_input.text()}")
        doc.add_paragraph("Початкові дані:")
        doc.add_table(rows=1, cols=len(self.df.columns))
        for r in self.df.values.tolist():
            row = doc.add_table(rows=1, cols=len(r)).rows[0].cells
            for i, val in enumerate(r):
                row[i].text = str(val)
        doc.add_paragraph("Результати аналізу:")
        for res in results:
            doc.add_paragraph(res)
        doc.add_paragraph(f"Дата аналізу: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph("Програма: SAD – Статистичний аналіз даних")
        doc.save("results.docx")

        QMessageBox.information(self, "Готово", "Аналіз завершено. Результати збережено у results.docx")
